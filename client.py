import socket
import tkinter as tk
from tkinter import filedialog, messagebox
import os
from docx import Document
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io

HOST = '127.0.0.1'
PORT = 65432

def send_string_to_server(message):
    try:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.connect((HOST, PORT))
            s.sendall(message.encode())
            data = s.recv(10000)
            return data.decode()
    except Exception as e:
        return f"Error: {e}"

def reverse_input_string():
    user_input = entry.get("1.0", tk.END).strip()
    if not user_input:
        messagebox.showwarning("Input Required", "Please enter a string.")
        return
    reversed_str = send_string_to_server(user_input)
    result_label.config(state='normal')  # Enable editing temporarily
    result_label.delete("1.0", tk.END)
    result_label.insert("1.0", reversed_str)
    result_label.config(state='disabled')  # Make read-only again

def load_file():
    filepath = filedialog.askopenfilename(
        title="Select a File",
        filetypes=[
            ("Text files", "*.txt"),
            ("Word files", "*.docx"),
            ("PDF files", "*.pdf"),
            ("All files", "*.*")
        ]
    )
    if not filepath:
        messagebox.showinfo("Info", "No file was selected")
        return

    try:
        content = ""
        file_extension = os.path.splitext(filepath)[1].lower()

        if file_extension == '.txt':
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read().strip()
        
        elif file_extension == '.docx':
            doc = Document(filepath)
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        elif file_extension == '.pdf':
            pdf = PdfReader(filepath)
            content_parts = []
            for page in pdf.pages:
                content_parts.append(page.extract_text())
            content = '\n'.join(content_parts)
        
        if not content:
            messagebox.showwarning("Warning", "The selected file is empty")
            return

        entry.delete("1.0", tk.END)
        entry.insert("1.0", content)
        messagebox.showinfo("Success", f"File loaded successfully: {filepath}")
    
    except Exception as e:
        messagebox.showerror("Error", f"Failed to read file: {str(e)}")

def save_reversed_text():
    result_label.config(state='normal')  # Enable temporarily to get text
    reversed_text = result_label.get("1.0", tk.END).strip()
    result_label.config(state='disabled')  # Make read-only again
    if not reversed_text or reversed_text.isspace():
        messagebox.showwarning("Warning", "No reversed text to save!")
        return
    
    filepath = filedialog.asksaveasfilename(
        title="Save Reversed Text",
        filetypes=[
            ("Text files", "*.txt"),
            ("Word files", "*.docx"),
            ("PDF files", "*.pdf"),
            ("All files", "*.*")
        ],
        defaultextension=".txt"
    )
    
    if not filepath:
        return
    
    try:
        file_extension = os.path.splitext(filepath)[1].lower()
        
        if file_extension == '.docx':
            doc = Document()
            doc.add_paragraph(reversed_text)
            doc.save(filepath)
        elif file_extension == '.pdf':
            c = canvas.Canvas(filepath, pagesize=letter)
            # Start writing at top of page with some margin
            y = letter[1] - 50  # Start 50 points from top
            # Split text into lines if it's too long
            lines = reversed_text.split('\n')
            for line in lines:
                words = line.split()
                current_line = ''
                for word in words:
                    if c.stringWidth(current_line + ' ' + word) < letter[0] - 100:  # Leave 50pt margin on each side
                        current_line += ' ' + word if current_line else word
                    else:
                        c.drawString(50, y, current_line)
                        y -= 20  # Move down 20 points
                        current_line = word
                    if y < 50:  # If near bottom of page
                        c.showPage()  # Start new page
                        y = letter[1] - 50  # Reset to top
                if current_line:  # Draw any remaining text
                    c.drawString(50, y, current_line)
                    y -= 20
            c.save()
        else:  # Default to txt
            with open(filepath, 'w', encoding='utf-8') as file:
                file.write(reversed_text)
        
        messagebox.showinfo("Success", f"File saved successfully as {filepath}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save file: {str(e)}")

def clear_text():
    entry.delete("1.0", tk.END)
    result_label.config(state='normal')
    result_label.delete("1.0", tk.END)
    result_label.config(state='disabled')

# GUI Setup
root = tk.Tk()
root.title("String Reversal Client")

# Set window size and position it in center
window_width = 600
window_height = 400
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
center_x = int(screen_width/2 - window_width/2)
center_y = int(screen_height/2 - window_height/2)
root.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')
root.configure(bg='#f0f0f0')  # Light gray background

# Connection Info Frame
conn_frame = tk.Frame(root, bg='#f0f0f0', padx=10, pady=5)
conn_frame.pack(fill=tk.X)
tk.Label(
    conn_frame,
    text=f"Host: {HOST}",
    font=('Helvetica', 10),
    bg='#f0f0f0',
    fg='#666666'
).pack(side=tk.LEFT, padx=(0, 15))
tk.Label(
    conn_frame,
    text=f"Port: {PORT}",
    font=('Helvetica', 10),
    bg='#f0f0f0',
    fg='#666666'
).pack(side=tk.LEFT)

# Title Frame
title_frame = tk.Frame(root, bg='#2c3e50', pady=15)  # Dark blue background
title_frame.pack(fill=tk.X)

title_label = tk.Label(
    title_frame,
    text="String Reversal Application",
    font=('Helvetica', 16, 'bold'),
    bg='#2c3e50',
    fg='white'
)
title_label.pack()

# Main Content Frame
content_frame = tk.Frame(root, bg='#f0f0f0', padx=20)
content_frame.pack(expand=True, fill=tk.BOTH, pady=20)

# Input Frame
input_frame = tk.Frame(content_frame, bg='#f0f0f0')
input_frame.pack(fill=tk.X, pady=(0, 10))

tk.Label(
    input_frame,
    text="Enter text or load from file:",
    font=('Helvetica', 14, 'bold'),
    bg='#f0f0f0'
).pack(pady=(0, 5))

# Create Text widget with scrollbar
text_frame = tk.Frame(input_frame, bg='#f0f0f0')
text_frame.pack(fill=tk.BOTH, expand=True)

entry = tk.Text(
    text_frame,
    font=('Helvetica', 14),
    height=5,
    wrap=tk.WORD,
    padx=10,
    pady=10
)
entry.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# Add scrollbar
scrollbar = tk.Scrollbar(text_frame)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

# Connect scrollbar to text widget
entry.config(yscrollcommand=scrollbar.set)
scrollbar.config(command=entry.yview)

# Button Frame
btn_frame = tk.Frame(content_frame, bg='#f0f0f0')
btn_frame.pack(pady=10)

button_style = {
    'font': ('Helvetica', 10),
    'width': 15,
    'pady': 5,
    'cursor': 'hand2'
}

reverse_btn = tk.Button(
    btn_frame,
    text="Reverse String",
    command=reverse_input_string,
    bg='#ADD8E6',  # Light blue
    fg='black',
    **button_style
)
reverse_btn.pack(side=tk.LEFT, padx=5)

load_btn = tk.Button(
    btn_frame,
    text="Load from File",
    command=load_file,
    bg='#ADD8E6',  # Light blue
    fg='black',
    **button_style
)
load_btn.pack(side=tk.LEFT, padx=5)

clear_btn = tk.Button(
    btn_frame,
    text="Clear",
    command=clear_text,
    bg='#808080',  # Grey
    fg='white',
    **button_style
)
clear_btn.pack(side=tk.LEFT, padx=5)

# Result Frame
result_frame = tk.Frame(content_frame, bg='#f0f0f0')
result_frame.pack(fill=tk.X, pady=20)

# Label for "Reversed:" text
tk.Label(
    result_frame,
    text="Reversed String:",
    font=('Helvetica', 14, 'bold'),
    bg='#f0f0f0'
).pack(pady=(0, 5))

# Create Text widget with scrollbar for result
result_text_frame = tk.Frame(result_frame, bg='#f0f0f0')
result_text_frame.pack(fill=tk.BOTH, expand=True)

result_label = tk.Text(
    result_text_frame,
    font=('Helvetica', 14),
    bg='#ffffff',  # White background
    relief=tk.SUNKEN,  # Add a sunken border effect
    height=5,
    padx=10,
    pady=10,
    wrap=tk.WORD
)
result_label.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# Add scrollbar for result
result_scrollbar = tk.Scrollbar(result_text_frame)
result_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

# Connect scrollbar to text widget
result_label.config(yscrollcommand=result_scrollbar.set)
result_scrollbar.config(command=result_label.yview)

# Make the result text widget read-only
result_label.config(state='disabled')

# Add save button under the result text area
save_btn = tk.Button(
    result_frame,
    text="Save Reversed Text",
    command=save_reversed_text,
    bg='#ADD8E6',  # Light blue
    fg='black',
    **button_style
)
save_btn.pack(pady=(10, 0))

# Add hover effects for buttons
def on_enter(e):
    e.widget['bg'] = {
        'Reverse String': '#87CEEB',  # Slightly darker light blue
        'Load from File': '#87CEEB',  # Slightly darker light blue
        'Save Reversed Text': '#87CEEB',  # Slightly darker light blue
        'Clear': '#696969'  # Darker grey
    }[e.widget['text']]

def on_leave(e):
    e.widget['bg'] = {
        'Reverse String': '#ADD8E6',  # Light blue
        'Load from File': '#ADD8E6',  # Light blue
        'Save Reversed Text': '#ADD8E6',  # Light blue
        'Clear': '#808080'  # Grey
    }[e.widget['text']]

for btn in [reverse_btn, load_btn, clear_btn, save_btn]:
    btn.bind("<Enter>", on_enter)
    btn.bind("<Leave>", on_leave)

root.mainloop()
